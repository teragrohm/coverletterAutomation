import sys
from time import gmtime, strftime, sleep
import re, pyperclip
from docx import Document
import webbrowser, requests
from bs4 import BeautifulSoup

coverLetter = Document('Cover_Letter.docx')


prev_company = coverLetter.paragraphs[6].text.replace(' ','')
run_num = 0

def iterate(list,arg):

    for k in range(len(list)):

        if arg == 'prev_job':

            if list[k].bold:

                #print(list[k].text)
                return k  

        elif arg == 'job_title':

            list[k] = list[k].capitalize()

        
def replaceJobTitle(title):

    first_paragraph = coverLetter.paragraphs[14].runs

    prev_job = iterate(first_paragraph,'prev_job')

    coverLetter.paragraphs[14].runs[prev_job].text = title

def replaceCompany(company):

    new_company = pyperclip.paste()

    #webbrowser.open('https://www.linkedin.com/jobs/search?keywords=' + new_company + '&location=Philippines&geoId=103121230&trk=public_jobs_jobs-search-bar_search-submit&currentJobId=3958150105&position=35&pageNum=0')

    i = 0

    for run in coverLetter.paragraphs[14].runs:

        line = run.text 
        line = line.replace(' ', '')

        company_name = re.search(company, line)

        if company_name:

            print(company_name.group())
            break

        i += 1

    coverLetter.paragraphs[14].runs[i].text = new_company
    return new_company

#coverLetter.paragraphs[6].text = replaceCompany(prev_company)

def replaceRecruiter():

    job = sys.argv[1:]

    #for i in range(len(job)):

     #    job[i] = job[i].capitalize()
    
    iterate(job,'job_title')
    job = ' '.join(job)

    replaceJobTitle(job)
    coverLetter.paragraphs[6].text = query = replaceCompany(prev_company)

    name = ''

    res = requests.get('https://www.linkedin.com/jobs/search?keywords=' + query + '&location=Philippines&geoId=103121230&trk=public_jobs_jobs-search-bar_search-submit&currentJobId=3958150105&position=35&pageNum=0')
    res.raise_for_status()

    soup = BeautifulSoup(res.content, 'html.parser')

    elems = soup.select('li .sr-only')


    for elem in elems:


        str_elem = elem.getText()
        str_elem = str_elem.replace('\n', '')

        str_match = re.search(job, str_elem)

        if str_match:

            job_link = elem.find_parent('a').get('href')

            print(str_match.group())

            sleep(3)
            
            job_post = requests.get(job_link)
            job_post.raise_for_status()

            soup = BeautifulSoup(job_post.content, 'html.parser')

            if soup.find('div', class_='base-main-card__info'):

              #webbrowser.open(job_link)
              
              poster_name = soup.select('.base-main-card__title')[0] 
              name = poster_name.getText()
              name = name.replace('\n','')
              name = re.findall(r'\w{2,}', name)
              print(name)
              
              salutation = coverLetter.paragraphs[11].runs

              salutation[2].text = input('Mr. or Ms.? ')
              salutation[2].text += ' '
              salutation[3].text = name[1]

              name = (' '.join(name))

              poster_pos = soup.select('.base-main-card__subtitle')[0]
              position = poster_pos.getText()
              position = position.replace('\n','')
              position = re.findall(r'\w{2,}', position)
              position = (' '.join(position))

              print('Job poster: ' + name)
              print('Position: ' + position)

              break
        
        sleep(3)

    webbrowser.open('https://www.google.com.ph/search?q=' + query)
    webbrowser.open('https://www.facebook.com/search/pages?q=' + query)
    
    return name, position


coverLetter.paragraphs[2].text = strftime("%B %d, %Y", gmtime())

recruiter_name, recruiter_position = replaceRecruiter()

coverLetter.paragraphs[4].text = recruiter_name

for run in coverLetter.paragraphs[4].runs:

 run.bold = True

coverLetter.paragraphs[5].text = recruiter_position

coverLetter.save('Cover_Letter.docx')